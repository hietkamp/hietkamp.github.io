#!/usr/bin/env python3
"""
Build Excel and Word templates for MVA work products.
Output: docs/wp/downloads/mva-*.xlsx / .docx
"""

from pathlib import Path
from openpyxl import Workbook
from openpyxl.styles import (
    Font, PatternFill, Alignment, Border, Side, numbers
)
from openpyxl.utils import get_column_letter
from openpyxl.formatting.rule import DataBarRule, FormulaRule
from openpyxl.worksheet.datavalidation import DataValidation
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

OUT = Path(__file__).parent.parent / "docs" / "wp" / "downloads"
OUT.mkdir(parents=True, exist_ok=True)

# ── Colour palette ────────────────────────────────────────────────────────────
PRAC    = "0F6E63"   # MVA green
INK     = "1C1A16"
PAPER   = "F3EFE6"
PAPER2  = "ECE6D8"
CARD    = "FBF9F3"
LINE    = "D8D0BF"
SOFT    = "524D42"
WHITE   = "FFFFFF"

MOT_BG  = "FDF5E0"   # motivation / driver
MOT_BD  = "E6C97A"
APP_BG  = "DCEEFF"   # application
APP_BD  = "8AB0E0"
TECH_BG = "D0EEDC"   # technology / paved road
TECH_BD = "7EC4A0"
IMP_BG  = "E8F0FA"   # implementation
IMP_BD  = "8AB0E0"
RED     = "B1492E"
AMBER   = "BB7A16"
GREY    = "8A93A2"

# ── Style helpers ─────────────────────────────────────────────────────────────
def fill(hex_): return PatternFill("solid", fgColor=hex_)
def font(bold=False, size=10, color=INK, italic=False, name="Calibri"):
    return Font(bold=bold, size=size, color=color, italic=italic, name=name)
def center(): return Alignment(horizontal="center", vertical="center", wrap_text=True)
def left():   return Alignment(horizontal="left",   vertical="top",    wrap_text=True)
def mid():    return Alignment(horizontal="left",   vertical="center", wrap_text=True)
def thin_border(sides="all"):
    s = Side(border_style="thin", color="C0B8A8")
    if sides == "all": return Border(left=s, right=s, top=s, bottom=s)
    b = Border()
    for side in sides: setattr(b, side, s)
    return b
def header_row(ws, row, cols, texts, bg=PRAC, fg=WHITE, row_h=18):
    for col, text in zip(cols, texts):
        c = ws.cell(row=row, column=col, value=text)
        c.font = font(bold=True, size=9, color=fg)
        c.fill = fill(bg)
        c.alignment = center()
        c.border = thin_border()
    ws.row_dimensions[row].height = row_h
def data_row(ws, row, cols, values, bg=None, row_h=None):
    for col, val in zip(cols, values):
        c = ws.cell(row=row, column=col, value=val)
        c.font = font(size=9)
        c.alignment = left()
        c.border = thin_border()
        if bg: c.fill = fill(bg)
    if row_h: ws.row_dimensions[row].height = row_h
def set_col_widths(ws, widths):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

# ── Word helpers ──────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1, color=PRAC):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_kicker(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(PRAC)
    run.font.bold = True
    p.paragraph_format.space_after = Pt(2)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string(SOFT)
    # left border via XML shading — use a tinted background instead
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'E6F4EC')
    p._p.get_or_add_pPr().append(shading)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(8)
    return p

def meta_table(doc, rows):
    table = doc.add_table(rows=len(rows)+1, cols=2)
    table.style = "Table Grid"
    hc = (PRAC, WHITE)
    for i, label in enumerate(["Veld", "Waarde"]):
        c = table.cell(0, i)
        c.text = label
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(hc[i] if i==1 else WHITE)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:fill'), PRAC)
        c._tc.get_or_add_tcPr().append(shading)
    for r, (label, placeholder) in enumerate(rows, 1):
        lc = table.cell(r, 0)
        lc.text = label
        lc.paragraphs[0].runs[0].font.size = Pt(9)
        lc.paragraphs[0].runs[0].font.bold = True
        vc = table.cell(r, 1)
        vc.text = placeholder
        vc.paragraphs[0].runs[0].font.size = Pt(9)
        vc.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(GREY)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:fill'), PAPER2.replace("#",""))
        lc._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()
    return table

def section_table(doc, headers, rows, col_widths=None, header_bg=PRAC):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.cell(0, i)
        c.text = h
        run = c.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(WHITE)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:fill'), header_bg)
        c._tc.get_or_add_tcPr().append(shading)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r, row_data in enumerate(rows, 1):
        bg = PAPER if r % 2 == 0 else CARD
        for i, val in enumerate(row_data):
            c = table.cell(r, i)
            c.text = str(val)
            run = c.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(SOFT if val and val.startswith("[") else INK)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:fill'), bg.replace("#",""))
            c._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def placeholder_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(GREY)
    run.font.italic = True
    p.paragraph_format.space_after = Pt(6)
    return p

def set_doc_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.0):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)

def add_title_block(doc, kicker, title, subtitle=None):
    add_kicker(doc, kicker)
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(INK)
        run.font.size = Pt(20)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor.from_string(PRAC)
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# WORD TEMPLATES
# ═══════════════════════════════════════════════════════════════════════════════

def build_word_architecture_vision():
    doc = Document()
    set_doc_margins(doc)
    add_title_block(doc,
        "MVA Work Product · ArchiMate Motivatielaag",
        "Architecture Vision",
        "de kleinste gedeelde intentie die verkeerde keuzes voorkomt")

    add_note(doc,
        "Dit document beschrijft de doelrichting, scope en waarde op één pagina. "
        "In COTS-context: de eerste deliverable vóór enige vendor geselecteerd wordt. "
        "Raakt alle ArchiMate-lagen op richtinggevend niveau: Motivatie (drivers, doelen, "
        "principes, uitsluitingscriteria) → Bedrijfslaag (scope) → Applicatielaag "
        "(COTS-richting, integratiepatroon) → Technologielaag (platform, paved road).")

    meta_table(doc, [
        ("Systeem / project", "[naam van het systeem of project]"),
        ("Versie",            "1.0"),
        ("Datum",             "[JJJJ-MM-DD]"),
        ("Eigenaar",          "[Lead/Enterprise Architect]"),
        ("Status",            "Concept | Overeengekomen | Gepubliceerd"),
        ("Goedgekeurd door",  "[naam sponsorende stakeholder]"),
    ])

    # Value statement
    add_heading(doc, "1. Waardeverklaring", 2)
    add_note(doc, "Formuleer in één alinea voor een manager, niet voor een architect. "
             "Structuur: Voor [doelgroep] die [behoefte] heeft, biedt dit systeem [kernoplossing] "
             "die, anders dan [alternatief], [onderscheidend voordeel] levert.")
    placeholder_para(doc, "[Waardeverklaring — max. 3 zinnen, geen jargon]")
    doc.add_paragraph()

    # Scope
    add_heading(doc, "2. Scope", 2)
    section_table(doc,
        ["Dimensie", "Binnen scope", "Buiten scope"],
        [
            ["Organisatie",  "[welke afdelingen/entiteiten]", "[expliciet uitgesloten]"],
            ["Processen",    "[welke primaire processen]",    "[buiten scope processen]"],
            ["Systemen",     "[welke systemen]",              "[niet in scope systemen]"],
            ["Data",         "[welke datagebieden]",          "[niet in scope data]"],
            ["Tijdhorizon",  "[periode en fases]",            "[wat volgt later]"],
        ],
        col_widths=[4, 7, 7]
    )

    # Drivers
    add_heading(doc, "3. Architecturele drivers (top 5)", 2)
    add_note(doc,
        "Formuleer elk driver falsifieerbaar: niet 'verbeter efficiëntie' maar "
        "'reduceer verwerkingstijd van 48 naar 6 uur'. "
        "Deze criteria zijn tevens de selectiecriteria voor COTS-componenten. "
        "Types: Strategisch doel · Kwaliteitsattribuut · Beperking · Principe.")
    section_table(doc,
        ["#", "Driver", "Type", "Falsifieerbaar criterium (meting)", "Prioriteit"],
        [
            ["D1", "[naam driver]", "[type]", "[criterium + meting + drempelwaarde]", "Hoog"],
            ["D2", "[naam driver]", "[type]", "[criterium + meting + drempelwaarde]", "Hoog"],
            ["D3", "[naam driver]", "[type]", "[criterium + meting + drempelwaarde]", "Middel"],
            ["D4", "[naam driver]", "[type]", "[criterium + meting + drempelwaarde]", "Middel"],
            ["D5", "[naam driver]", "[type]", "[criterium + meting + drempelwaarde]", "Laag"],
        ],
        col_widths=[1, 4.5, 3.5, 7, 2.5]
    )

    # Principles
    add_heading(doc, "4. Leidende architectuurprincipes", 2)
    add_note(doc, "Principes zijn keuzeregels, geen waarden. Een principe helpt een dilemma "
             "te beslechten — als het altijd waar is, is het geen principe maar een feit. "
             "Gebruik principes actief om componentopties af te wijzen.")
    section_table(doc,
        ["P#", "Principe", "Consequentie voor ontwerp"],
        [
            ["P1", "Configureer, customiseer niet",     "COTS-aanpassingen buiten vendor-API vereisen een ADR"],
            ["P2", "Standaarden boven maatwerk",         "Waar een open standaard bestaat, is die leidend"],
            ["P3", "Exitstrategie vóór adoptie",         "Elk COTS-product heeft een gedocumenteerde migratieroute"],
            ["P4", "Interoperabiliteit by design",       "Elk systeem exposeert een gedocumenteerde API"],
            ["P5", "[eigen principe]",                   "[consequentie]"],
        ],
        col_widths=[1.5, 7, 9]
    )

    # Exclusion criteria
    add_heading(doc, "5. Uitsluitingscriteria", 2)
    add_note(doc,
        "Technologieën, leveranciers, architectuurpatronen of licentiemodellen die expliciet "
        "buiten beschouwing worden gelaten. In COTS essentieel: maakt verkeerde vendor-selecties "
        "onmogelijk vóórdat er ook maar één demo plaatsvindt.")
    section_table(doc,
        ["Categorie", "Uitgesloten", "Reden"],
        [
            ["Leverancier",        "[naam leverancier/product]", "[reden: lock-in / compliance / licentiemodel]"],
            ["Architectuurpatroon","[patroon, bijv. point-to-point DB]", "[reden]"],
            ["Licentiemodel",      "[bijv. per-CPU perpetueel]",         "[reden]"],
            ["",                   "",                                    ""],
        ],
        col_widths=[3.5, 6, 9]
    )

    # Architecture direction
    add_heading(doc, "6. Architectuurrichting", 2)
    add_note(doc,
        "Eén zin per ArchiMate-laag: welk platform, welk integratiepatroon, welke paved road. "
        "Geen details — die volgen in de Architecture Definition.")
    section_table(doc,
        ["Laag", "Richting (één zin)"],
        [
            ["Bedrijfslaag",    "[welke processen worden ondersteund]"],
            ["Applicatielaag",  "[COTS-selectierichting en integratiepatroon]"],
            ["Technologielaag", "[platform en paved road-richting]"],
        ],
        col_widths=[4, 14]
    )

    # Irreversible decisions
    add_heading(doc, "7. Signalering onomkeerbare beslissingen", 2)
    add_note(doc,
        "Benoem welke beslissingen in dit initiatief kostbaar zijn om terug te draaien. "
        "Dit zijn aandachtspunten voor ADR's — ook als ze nog niet genomen zijn. "
        "Focus op: vendor lock-in, datamodel, integratie-API, licentiemodel.")
    section_table(doc,
        ["Beslissing (nog te nemen)", "Categorie", "Risico bij terugdraaien"],
        [
            ["[bijv. keuze cloudplatform]", "Vendor lock-in",  "[hoog / middel / laag — toelichting]"],
            ["[bijv. datamodelkeuze]",      "Datamodel",        "[toelichting]"],
            ["[bijv. SSO-protocol]",        "Integratie-API",   "[toelichting]"],
        ],
        col_widths=[6, 3.5, 9]
    )

    # Checklist
    add_heading(doc, "8. Controle vóór publicatie", 2)
    checks = [
        "Waardeverklaring is begrijpelijk voor een niet-technisch publiek",
        "Scope is expliciet: in én buiten scope benoemd",
        "Elke driver heeft een falsifieerbaar succescriterium",
        "Uitsluitingscriteria zijn benoemd — verkeerde vendor-selecties zijn onmogelijk",
        "Architectuurrichting is per laag beschreven en herleidbaar naar drivers",
        "Onomkeerbare beslissingen zijn gesignaleerd als aandachtspunten",
        "Principes zijn overeengekomen met de sponsorende stakeholder",
        "Sponsorende stakeholder heeft document formeel geaccordeerd",
        "Document past op één A3 of twee A4-pagina's",
    ]
    for chk in checks:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"☐  {chk}")
        run.font.size = Pt(9)

    doc.save(OUT / "mva-architecture-vision.docx")
    print("✓  mva-architecture-vision.docx")


def build_word_architecture_definition():
    doc = Document()
    set_doc_margins(doc)
    add_title_block(doc,
        "MVA Work Product · ArchiMate Applicatie- & Technologielaag",
        "Architecture Definition",
        "het Minimum Viable Architecture (MVA)")

    add_note(doc,
        "Bevat uitsluitend de bouwstenen die een huidige beslissing vereist. "
        "Doorgestreepte of uitgestelde onderdelen worden expliciet benoemd. "
        "ArchiMate-lagen: Bedrijf → Applicatie (COTS) → Technologie (paved road platform).")

    meta_table(doc, [
        ("Systeem / project", "[naam]"),
        ("Versie",            "0.1"),
        ("Datum",             "[JJJJ-MM-DD]"),
        ("Architect",         "[naam]"),
        ("Status",            "Skelet | Werkend | Stabiel"),
        ("Gebaseerd op ADR's","[ADR-01, ADR-02, ...]"),
    ])

    # Business layer
    add_heading(doc, "1. Bedrijfslaag — processen & gebruikers", 2)
    add_note(doc, "Welke bedrijfsprocessen en -rollen worden ondersteund? "
             "ArchiMate: Business Process, Business Function, Business Role.")
    section_table(doc,
        ["Bedrijfsproces / -functie", "Beschrijving", "Gebruikersrol"],
        [
            ["[Primair proces]",      "[beschrijving]", "[rol]"],
            ["[Ondersteunend proces]","[beschrijving]", "[rol]"],
            ["[Beheer]",             "[beschrijving]", "[admin-rol]"],
        ],
        col_widths=[5, 9, 4.5]
    )

    # Application layer
    add_heading(doc, "2. Applicatielaag — COTS-componenten", 2)
    add_note(doc,
        "Eén rij per significant COTS-product. Markeer paved-road producten. "
        "ArchiMate: Application Component (COTS), Application Service, Application Interface.")
    section_table(doc,
        ["Component", "Leverancier", "Versie / Tier", "Paved Road", "Dekt ADR"],
        [
            ["[COTS Pakket A]",  "[leverancier]", "SaaS Enterprise", "Ja", "ADR-02"],
            ["[COTS Pakket B]",  "[leverancier]", "SaaS",            "Ja", "ADR-03"],
            ["[Integratie-laag]","[leverancier]", "Managed",         "Ja", "ADR-04"],
            ["[Portaal / UI]",   "[leverancier]", "SaaS",            "Ja", "ADR-03"],
            ["[Component X]",    "[leverancier]", "[tier]",          "Nee → ADR-0x", "—"],
        ],
        col_widths=[4.5, 3.5, 3.5, 2.5, 3]
    )

    # Integration layer — first class
    add_heading(doc, "3. Integratielaag (glueware) — first-class element", 2)
    add_note(doc,
        "De integratielaag is altijd first-class — nooit een implementatiedetail achteraf. "
        "Inadequate integratieplanning veroorzaakt 42% van alle COTS-implementatieproblemen. "
        "Documenteer per koppeling: eigenaar van de glueware, adaptatiekosten en "
        "betrouwbaarheidsimpact. ArchiMate: Application Interface + Triggering/Flow Relationship.")
    section_table(doc,
        ["Van", "Naar", "Integratiepatroon", "Protocol", "Glueware-eigenaar", "Adaptatiekosten", "ADR"],
        [
            ["[Pakket A]", "[Integratie-laag]", "Async event",    "CloudEvents", "[team]", "Laag",   "ADR-04"],
            ["[Pakket B]", "[Integratie-laag]", "REST",           "OpenAPI 3.x", "[team]", "Laag",   "ADR-04"],
            ["[Portaal]",  "[Pakket A]",        "API-aanroep",    "REST/OIDC",   "[team]", "Middel", "ADR-03"],
            ["[Legacy X]", "[Integratie-laag]", "SOAP/XML",       "SOAP 1.2",    "[team]", "Hoog — tijdelijk, ADR vereist", "ADR-0x"],
        ],
        col_widths=[2.5, 2.5, 3, 2.5, 2.5, 3.5, 2]
    )

    # Technology layer
    add_heading(doc, "4. Technologielaag — paved road platform", 2)
    add_note(doc,
        "Infrastructuur en platforms waarop de applicatiecomponenten draaien. "
        "ArchiMate: Technology Node, System Software, Technology Service.")
    section_table(doc,
        ["Platform / node", "Categorie", "Leverancier", "Regio / Zone", "Standaard"],
        [
            ["[Cloud Platform]",    "Cloud",       "[leverancier]", "EU / NL",     "Paved Road"],
            ["[Identity Provider]", "Identiteit",  "[leverancier]", "Tenant: [x]", "SSO verplicht"],
            ["[CI/CD Platform]",    "Delivery",    "[leverancier]", "SaaS",        "GitOps-patroon"],
            ["[Monitoring]",        "Observability","[leverancier]", "SaaS",        "Paved Road"],
        ],
        col_widths=[4.5, 3, 3, 3.5, 3.5]
    )

    # Deferred
    add_heading(doc, "5. Bewust uitgesteld (buiten scope MVA)", 2)
    add_note(doc, "Explicitering van wat NIET in dit document zit en waarom. "
             "Dit maakt het minimale karakter van de MVA zichtbaar.")
    section_table(doc,
        ["Onderdeel", "Reden van uitstel", "Moment van heroverweging"],
        [
            ["Uitputtend componentmodel", "Niet nodig voor huidige beslissingen", "Na ADR-05"],
            ["Volledig datamodel",        "Wacht op productkeuze (ADR-02)",        "H2"],
            ["Detailarchitectuur [X]",   "[reden]",                               "[trigger]"],
        ],
        col_widths=[5, 8, 5]
    )

    doc.save(OUT / "mva-architecture-definition.docx")
    print("✓  mva-architecture-definition.docx")


def build_word_adr():
    doc = Document()
    set_doc_margins(doc)
    add_title_block(doc,
        "MVA Work Product · Architecture Decision Record",
        "ADR-[NUMMER]: [Beslissingstitel]",
        "een beredeneerde keuze")

    add_note(doc,
        "Eén ADR per significante, kostbaar-om-terug-te-draaien beslissing (Nygard 2011). "
        "De ADR is nooit de Paved Road zelf — hij stelt een Paved Road-element in, "
        "verantwoordt een off-road-afwijking, of herziet een bestaande standaard. "
        "Leg de redenering vast — niet alleen de uitkomst.")

    # Y-Statement summary
    add_heading(doc, "Y-Statement (kernformulering)", 2)
    add_note(doc,
        "Formuleer de beslissing in één zin vóórdat je de rest invult. "
        "Als dit niet lukt, is de beslissing nog niet scherp genoeg. "
        "Structuur: In de context van [use case], geconfronteerd met [zorg], "
        "hebben we gekozen voor [optie] om [kwaliteit] te bereiken, "
        "waarbij we [nadeel] accepteren.")
    placeholder_para(doc, "[Y-Statement — één zin]")
    doc.add_paragraph()

    meta_table(doc, [
        ("ADR-nummer",                "ADR-[xx]"),
        ("Datum",                     "[JJJJ-MM-DD]"),
        ("Status",                    "Proposed | Accepted | Rejected | Deprecated | Superseded by ADR-xx"),
        ("Eigenaar",                  "[Lead/Enterprise Architect]"),
        ("Beslissers",                "[namen en rollen]"),
        ("Onomkeerbaarheidscategorie","Vendor lock-in | Datamodel | Integratie-API | Licentiemodel | Overig"),
        ("Paved Road-relatie",        "Stelt SR in | Herziet SR-[xx] | Off-road afwijking van SR-[xx] | Geen"),
        ("Laatste verantwoord moment","[JJJJ-MM-DD of trigger]"),
    ])

    add_heading(doc, "1. Context", 2)
    add_note(doc,
        "Beschrijf WHY, niet WHAT. Welke situatie maakt deze beslissing urgent? "
        "Welke krachten spelen mee, welke beperkingen gelden, waarom nu? "
        "Bevat ook: is dit een onomkeerbare beslissing en waarom?")
    section_table(doc,
        ["Gerelateerde driver(s)", "Gerelateerde ASR('s)", "Gerelateerde ADR's"],
        [["[D1, D3]", "[ASR-01, ASR-02]", "[ADR-02 (afhankelijk van deze)]"]],
        col_widths=[5, 5, 7.5]
    )
    placeholder_para(doc,
        "[Beschrijving van de context — 3–6 zinnen. Wat is er aan de hand? "
        "Welke keuze moet worden gemaakt en waarom nu? "
        "Is dit een onomkeerbare beslissing (vendor lock-in / datamodel / integratie-API)?]")

    add_heading(doc, "2. Overwogen opties", 2)
    add_note(doc,
        "Minimaal drie opties, inclusief de nul-optie (niets doen). "
        "Normaliseer criteria naar hetzelfde abstractieniveau — vergelijk niet een technisch "
        "voordeel van A met een organisatorisch nadeel van B. "
        "Voeg voor COTS-beslissingen de glueware-kosten per optie toe.")
    section_table(doc,
        ["Optie", "Omschrijving", "Voordelen", "Nadelen / risico's", "Glueware-kosten"],
        [
            ["A (gekozen)",   "[omschrijving]",         "[voordelen]",  "[risico's]",             "Laag / Middel / Hoog"],
            ["B",             "[alternatief]",           "[voordelen]",  "[nadelen]",              "Laag / Middel / Hoog"],
            ["C (nul-optie)", "Status quo handhaven",   "Geen kosten",  "Driver niet geadresseerd","—"],
        ],
        col_widths=[2.5, 5, 4, 4, 3]
    )

    add_heading(doc, "3. Beslissing", 2)
    placeholder_para(doc,
        "[Wij kiezen optie A omdat ... Geef de volledige redenering. "
        "Vermeld trade-offs die bewust worden geaccepteerd. "
        "Verwijzing naar de drivers en principes die de keuze sturen.]")

    add_heading(doc, "4. Consequenties", 2)
    add_note(doc, "In COTS: benoem expliciet de integratieafhankelijkheden die deze beslissing creëert.")
    section_table(doc,
        ["Positieve consequenties", "Te beheren consequenties"],
        [
            ["• [voordeel / oplossing van probleem]\n• [voordeel 2]",
             "• [risico / nieuwe afhankelijkheid]\n• Integratieafhankelijkheid: [beschrijving]\n• Beheersmaatregel: [actie]"],
        ],
        col_widths=[8.75, 8.75]
    )

    add_heading(doc, "5. Falsifieerbaarheidscriterium", 2)
    add_note(doc,
        "Hoe en wanneer weten we dat deze beslissing juist was? "
        "Formuleer als meetbaar criterium — het ASR dat door deze beslissing wordt vervuld. "
        "Dit is de basis voor validatie in Prove the Risky Parts.")
    placeholder_para(doc,
        "[Wanneer [meting X] uitkomst Y toont, is deze beslissing gevalideerd. "
        "Uitvoerder: [naam]. Streefdatum: [datum].]")
    section_table(doc,
        ["Falsifieerbaarheidscriterium", "Uitvoerder", "Streefdatum", "Status"],
        [
            ["[meetbaar criterium + drempelwaarde]", "[naam]", "[datum]", "Open | Lopend | Afgerond"],
        ],
        col_widths=[7.5, 3, 3, 4]
    )

    add_heading(doc, "6. Verwijzingen", 2)
    section_table(doc,
        ["Type", "Verwijzing"],
        [
            ["Driver",           "D1 · [naam]"],
            ["ASR",              "ASR-01 · [naam kwaliteitsattribuut]"],
            ["ADR",              "ADR-02 (verwant) · ADR-05 (superseded door)"],
            ["Paved Road SR",    "SR-[xx] — dit ADR stelt/herziet dit Standaard Record"],
            ["Work product",     "Paved Road catalogus · Architecture Definition"],
            ["Extern",           "[leverancierscontract / documentatie]"],
        ],
        col_widths=[4, 13.5]
    )

    doc.save(OUT / "mva-adr.docx")
    print("✓  mva-adr.docx")


def build_word_review_record():
    doc = Document()
    set_doc_margins(doc)
    add_title_block(doc,
        "MVA Work Product · Architecture Review Record (ARR)",
        "Architecture Review Record",
        "Conformiteitsrecord (CR) | Exception Request (ER)")

    add_note(doc,
        "De ARR beantwoordt de vraag: voldoet dit aan wat we hebben afgesproken? "
        "Niet: waarom hebben we dit gekozen — dat is het ADR. "
        "Twee uitkomsten: CR (on-road bevestigd) of ER (off-road geformaliseerd). "
        "Vul eerst het recordtype in — dat bepaalt welke secties van toepassing zijn.")

    meta_table(doc, [
        ("ARR-nummer",    "CR-[xx] | ER-[xx]"),
        ("Recordtype",    "Conformiteitsrecord (CR) | Exception Request (ER)"),
        ("Datum",         "[JJJJ-MM-DD]"),
        ("Aanvrager",     "[team / naam]"),
        ("Project",       "[projectnaam]"),
        ("Reviewer",      "[Lead Architect]"),
        ("Deelnemers",    "[namen en rollen]"),
        ("Uitkomst",      "CR: Conform | Conform met aandachtspunten  //  ER: ✔ Goedgekeurd | ⚠ Met condities | ✗ Afgekeurd"),
        ("Geldig tot",    "[datum herziening]"),
    ])

    add_heading(doc, "1. Aanleiding & beschrijving", 2)
    add_note(doc,
        "CR: beschrijf welk systeem/product/leverancier wordt getoetst. "
        "ER: beschrijf welk product of patroon van de Paved Road afwijkt en waarom "
        "de afwijking zakelijk gerechtvaardigd is.")
    placeholder_para(doc,
        "[Beschrijf waarvoor de review is aangevraagd. "
        "CR: leverancier X wordt getoetst aan Laag 2-kwalificatiecriteria. "
        "ER: product Y wijkt af van SR-[xx] omdat ...]")

    add_heading(doc, "2. Beoordeeld element", 2)
    section_table(doc,
        ["Element", "Type", "Getoetste Standaard Records"],
        [
            ["[Beoordeeld product/patroon/leverancier]",
             "[App Component / Tech Node / Leverancier]",
             "SR-[xx]: [naam standaard]"],
        ],
        col_widths=[5, 4, 8.5]
    )

    # ER-only section
    add_heading(doc, "3. Afwijking van standaard (alleen ER)", 2)
    add_note(doc, "Alleen invullen bij Exception Request. Sla over bij Conformiteitsrecord.")
    section_table(doc,
        ["Standaard Record", "Afwijking", "Specificiek aspect"],
        [
            ["SR-[xx] · [naam]", "[wat wijkt af]", "[welk criterium wordt niet gehaald]"],
        ],
        col_widths=[4.5, 7, 6]
    )

    add_heading(doc, "4. Toetsing", 2)
    add_note(doc,
        "Toets altijd aan drivers ÉN principes ÉN Paved Road Standaard Records. "
        "✔ = Voldoet · ⚠ = Voldoet met kanttekening · ✗ = Voldoet niet")
    section_table(doc,
        ["Criterium (SR / Driver / Principe)", "Bevinding (feit, geen mening)", "Oordeel"],
        [
            ["SR-[xx] · [naam standaard]",        "[bevinding]", "✔ / ⚠ / ✗"],
            ["Driver D1 · [naam]",                "[bevinding]", "✔ / ⚠ / ✗"],
            ["Principe P1 · Configureer niet",    "[bevinding]", "✔ / ⚠ / ✗"],
            ["ASR-01 · [kwaliteitsattribuut]",    "[bevinding]", "✔ / ⚠ / ✗"],
            ["Kosten / TCO / glueware-kosten",    "[bevinding]", "✔ / ⚠ / ✗"],
        ],
        col_widths=[5.5, 9, 3]
    )

    add_heading(doc, "5. Compenserende maatregelen (alleen ER)", 2)
    add_note(doc,
        "Alleen invullen bij Exception Request. "
        "Welke maatregelen compenseren de risico's van de afwijking? "
        "Formuleer SMART: wat, wie, wanneer.")
    section_table(doc,
        ["Maatregel", "Eigenaar", "Streefdatum"],
        [
            ["[maatregel 1 — compenseert risico X]", "[naam]", "[datum]"],
            ["[maatregel 2]",                         "[naam]", "[datum]"],
        ],
        col_widths=[10, 4, 3.5]
    )

    add_heading(doc, "6. Tijdshorizon & eigenaarschap (alleen ER)", 2)
    add_note(doc,
        "Elke Exception Request is tijdelijk of permanent — maak dit expliciet. "
        "Zonder tijdshorizon wordt de uitzondering een permanente afwijking. "
        "De eigenaar is verantwoordelijk voor naleving en herziening.")
    section_table(doc,
        ["Veld", "Waarde"],
        [
            ["Tijdshorizon",           "Tijdelijk (vervalt [datum / trigger]) | Permanent (met motivatie)"],
            ["Eigenaar afwijking",     "[naam + rol — verantwoordelijk voor onderhoud]"],
            ["Herzieningsmoment",       "[plateau-overgang P[x] of vaste datum]"],
            ["Migratiepad",            "[pad terug naar on-road, of onderbouwing permanentie]"],
        ],
        col_widths=[5, 12.5]
    )

    add_heading(doc, "7. Uitkomst & condities", 2)
    placeholder_para(doc,
        "[Volledige redenering achter het oordeel. "
        "CR: beschrijf op welke punten voldaan wordt en eventuele aandachtspunten. "
        "ER: welke afweging is gemaakt, welke risico's worden geaccepteerd en door wie?]")
    section_table(doc,
        ["Conditie / actiepunt", "Eigenaar", "Streefdatum", "Status"],
        [
            ["CR: [aandachtspunt X] monitoren",                "[naam]", "[datum]", "Open"],
            ["ER: ADR-0x opstellen voor deze beslissing",      "[architect]", "[datum]", "Open"],
            ["ER: Paved Road Catalogus updaten (CR of status Afwijking)", "[architect]", "[datum]", "Open"],
            ["Bevinding rapporteren aan MVG-governance",        "[naam]", "[datum]", "Open"],
        ],
        col_widths=[8, 3, 3, 2.5]
    )

    add_heading(doc, "8. Handtekeningen", 2)
    section_table(doc,
        ["Rol", "Naam", "Handtekening", "Datum"],
        [
            ["Lead / Enterprise Architect", "[naam]", " " * 30, "[datum]"],
            ["Eigenaar afwijking (ER)",     "[naam]", " " * 30, "[datum]"],
            ["Aanvrager",                   "[naam]", " " * 30, "[datum]"],
        ],
        col_widths=[4.5, 4, 6, 3]
    )

    doc.save(OUT / "mva-review-record.docx")
    print("✓  mva-review-record.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# EXCEL TEMPLATES
# ═══════════════════════════════════════════════════════════════════════════════

def build_excel_asr():
    wb = Workbook()

    # ── Sheet 1: QAS-Scenarios ────────────────────────────────────────────────
    ws = wb.active
    ws.title = "QAS-Scenarios"
    ws.sheet_view.showGridLines = False

    # Title band
    ws.merge_cells("A1:H1")
    t = ws["A1"]
    t.value = "Architecturally-Significant Requirements (ASR's) — Kwaliteitsattribuutscenario's"
    t.font = font(bold=True, size=13, color=WHITE)
    t.fill = fill(PRAC)
    t.alignment = center()
    ws.row_dimensions[1].height = 26

    ws.merge_cells("A2:H2")
    s = ws["A2"]
    s.value = ("Documenteer uitsluitend requirements die significant genoeg zijn om de architectuur te vormen. "
               "Elk heeft een meetbaar succescriterium en is gekoppeld aan een driver.")
    s.font = font(size=8, italic=True, color=SOFT)
    s.fill = fill(PAPER2)
    s.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    ws.row_dimensions[2].height = 28

    # Headers row 3
    headers = ["ASR #", "Type", "Kwaliteitsattribuut / ASFR-categorie",
               "Scenario (stimulus → reactie)", "Meetbaar criterium",
               "Driver", "Prioriteit", "COTS-selectiecriterium", "ADR"]
    header_row(ws, 3, range(1, 10), headers, bg=PRAC)

    type_dv = DataValidation(
        type="list",
        formula1='"NFR (kwaliteitsattribuut),ASFR (functioneel met architectuurimpact)"',
        allow_blank=True
    )
    ws.add_data_validation(type_dv)

    qa_dv = DataValidation(
        type="list",
        formula1='"Beschikbaarheid,Performance,Beveiliging,Schaalbaarheid,Betrouwbaarheid,'
                 'Onderhoudbaarheid,Conformance,Interoperabiliteit,'
                 'Audit Trail,Batch Processing,Localization,Notificaties/Alerts,'
                 'Data-invoer,Betaling,Rapportage,Zoeken,Third Party Interaction,'
                 'Workflow,Licenties,Gebruikersgedraganalyse,Opslagmechanismen"',
        allow_blank=True
    )
    ws.add_data_validation(qa_dv)

    priority_dv = DataValidation(
        type="list", formula1='"Hoog,Middel,Laag"', allow_blank=True
    )
    ws.add_data_validation(priority_dv)

    cots_dv = DataValidation(
        type="list", formula1='"Ja — opnemen in RFP/PvE,Nee,Nader te bepalen"', allow_blank=True
    )
    ws.add_data_validation(cots_dv)

    sample = [
        ["ASR-01", "NFR (kwaliteitsattribuut)", "Beschikbaarheid",
         "Bij uitval van één node blijft het systeem operationeel. Failover < 30 s.",
         "99,5% uptime per kwartaal", "D1", "Hoog", "Ja — opnemen in RFP/PvE", "ADR-01"],
        ["ASR-02", "NFR (kwaliteitsattribuut)", "Performance",
         "Onder normaal gebruik (500 gelijktijdige gebruikers) reageert 95% van verzoeken < 2 s.",
         "p95 responstijd < 2 s", "D2", "Hoog", "Ja — opnemen in RFP/PvE", "ADR-04"],
        ["ASR-03", "NFR (kwaliteitsattribuut)", "Beveiliging",
         "Alle persoonsgegevens worden versleuteld opgeslagen en getransporteerd.",
         "ISO 27001 conformant", "D3", "Hoog", "Ja — opnemen in RFP/PvE", "ADR-02"],
        ["ASR-04", "ASFR (functioneel met architectuurimpact)", "Third Party Interaction",
         "Systeem exporteert aanvragen naar backoffice leverancier X.",
         "Koppeling werkt conform OpenAPI 3.x contract", "D2", "Hoog", "Ja — opnemen in RFP/PvE", "ADR-03"],
        ["ASR-05", "ASFR (functioneel met architectuurimpact)", "Notificaties/Alerts",
         "Bij statuswijziging ontvangt de betrokken actor een notificatie binnen 30 s.",
         "99% aflevering < 30 s (monitoring)", "D1", "Middel", "Ja — opnemen in RFP/PvE", "ADR-04"],
        ["ASR-06", "NFR (kwaliteitsattribuut)", "Conformance",
         "Implementaties op paved road vereisen geen architectuurreview.",
         ">90% on-road implementaties", "D4", "Middel", "Nee", "—"],
        ["ASR-0x", "", "", "", "", "", "", "", ""],
    ]

    priority_colors = {"Hoog": "FDE8E0", "Middel": "FDF5E0", "Laag": PAPER}

    for r, row in enumerate(sample, 4):
        bg = priority_colors.get(row[6], CARD) if row[6] else CARD
        for c, val in enumerate(row, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.font = font(size=9, bold=(c == 1))
            cell.fill = fill(bg)
            cell.border = thin_border()
            cell.alignment = left()
        ws.row_dimensions[r].height = 36
        type_dv.sqref += f"B{r}:B{r}"
        qa_dv.sqref += f"C{r}:C{r}"
        priority_dv.sqref += f"G{r}:G{r}"
        cots_dv.sqref += f"H{r}:H{r}"

    set_col_widths(ws, [8, 22, 22, 34, 22, 6, 9, 24, 9])
    ws.freeze_panes = "A4"
    ws.auto_filter.ref = f"A3:I{3+len(sample)}"

    ws.conditional_formatting.add(
        f"A4:I{3+len(sample)}",
        FormulaRule(formula=['$G4="Hoog"'], fill=PatternFill("solid", fgColor="FDE8E0"))
    )
    ws.conditional_formatting.add(
        f"A4:I{3+len(sample)}",
        FormulaRule(formula=['$G4="Middel"'], fill=PatternFill("solid", fgColor="FDF5E0"))
    )

    ws.print_area = f"A1:I{3+len(sample)}"
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1

    # ── Sheet 2: Beperkingen ─────────────────────────────────────────────────
    ws2 = wb.create_sheet("Beperkingen")
    ws2.sheet_view.showGridLines = False
    ws2.merge_cells("A1:E1")
    t2 = ws2["A1"]
    t2.value = "Beperkingen (Constraints)"
    t2.font = font(bold=True, size=13, color=WHITE)
    t2.fill = fill(PRAC)
    t2.alignment = center()
    ws2.row_dimensions[1].height = 24

    header_row(ws2, 2, range(1, 6),
               ["Beperking", "Type", "Beschrijving", "Impacteer ADR", "Eigenaar"])

    cons = [
        ["Datacenter NL-only",        "Regulatoir",       "AVG/BIO vereist opslag in NL of EU",                   "ADR-01", "[naam]"],
        ["Budget CAPEX jaar 1",        "Financieel",       "Max. investering jaar 1: [bedrag]",                     "ADR-05", "[naam]"],
        ["Bestaand IdP",               "Technisch",        "SSO via bestaande [Entra ID / Okta]",                  "ADR-02", "[naam]"],
        ["Integratie-API leverancier X","Technisch",       "Leverancier X ondersteunt alleen SOAP 1.2 — geen REST", "ADR-0x", "[naam]"],
        ["Licentiemodel beperking",    "Contractueel",     "Licentiemodel vendor Y is per-CPU — schaalbaarheidsbeperking", "ADR-0x", "[naam]"],
        ["",                           "",                 "",                                                      "",       ""],
    ]
    for r, row in enumerate(cons, 3):
        for c, val in enumerate(row, 1):
            cell = ws2.cell(row=r, column=c, value=val)
            cell.font = font(size=9)
            cell.fill = fill(PAPER if r % 2 == 0 else CARD)
            cell.border = thin_border()
            cell.alignment = left()
        ws2.row_dimensions[r].height = 22

    type_dv = DataValidation(
        type="list",
        formula1='"Regulatoir,Financieel,Technisch,Organisatorisch,Contractueel,Integratie-API,Licentiemodel,Tijdgebonden"',
        allow_blank=True
    )
    ws2.add_data_validation(type_dv)
    for r in range(3, 3+len(cons)):
        type_dv.sqref += f"B{r}:B{r}"

    set_col_widths(ws2, [22, 14, 36, 14, 14])
    ws2.freeze_panes = "A3"

    # ── Sheet 3: Drivers-referentie ──────────────────────────────────────────
    ws3 = wb.create_sheet("Drivers-referentie")
    ws3.sheet_view.showGridLines = False
    ws3.merge_cells("A1:F1")
    t3 = ws3["A1"]
    t3.value = "Architecturele Drivers — referentielijst voor ASR's"
    t3.font = font(bold=True, size=12, color=WHITE)
    t3.fill = fill(PRAC)
    t3.alignment = center()
    ws3.row_dimensions[1].height = 22

    header_row(ws3, 2, range(1, 7),
               ["Driver #", "Naam", "Type", "Eigenaar", "Prioriteit", "Status"])

    drv = [
        ["D1", "[Beschikbaarheid bedrijfsproces]", "Kwaliteitsattribuut", "[naam]", "Hoog",   "Gekwantificeerd"],
        ["D2", "[Performance eindgebruiker]",        "Kwaliteitsattribuut", "[naam]", "Hoog",   "Gekwantificeerd"],
        ["D3", "[Beveiliging & compliance]",          "Beperking",          "[naam]", "Hoog",   "Gekwantificeerd"],
        ["D4", "[Standaardisatie paved road]",        "Strategisch doel",   "[naam]", "Middel", "Geprioriteerd"],
        ["D5", "[Kosten TCO]",                        "Financieel",         "[naam]", "Middel", "Geïdentificeerd"],
        ["",   "",                                    "",                   "",       "",       ""],
    ]
    for r, row in enumerate(drv, 3):
        for c, val in enumerate(row, 1):
            cell = ws3.cell(row=r, column=c, value=val)
            cell.font = font(size=9, bold=(c == 1))
            cell.fill = fill(MOT_BG if row[0] else CARD)
            cell.border = thin_border()
            cell.alignment = left()
        ws3.row_dimensions[r].height = 20

    set_col_widths(ws3, [9, 28, 20, 14, 10, 16])
    ws3.freeze_panes = "A3"

    # ── Sheet 4: ASFR-categorieën referentie ────────────────────────────────────
    ws4 = wb.create_sheet("ASFR-categorieën")
    ws4.sheet_view.showGridLines = False
    ws4.merge_cells("A1:D1")
    t4 = ws4["A1"]
    t4.value = "ASFR-categorieën — functionele requirements met verborgen architecturele impact"
    t4.font = font(bold=True, size=12, color=WHITE)
    t4.fill = fill(PRAC)
    t4.alignment = center()
    ws4.row_dimensions[1].height = 24
    ws4.merge_cells("A2:D2")
    s4 = ws4["A2"]
    s4.value = ("Bron: Anish, Daneva, Cleland-Huang et al. (RE2015). "
                "Gebruik als checklist: elke herkenbare categorie verdient een Probing Question en mogelijk een eigen ASR.")
    s4.font = font(size=8, italic=True, color=SOFT)
    s4.fill = fill(PAPER2)
    s4.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    ws4.row_dimensions[2].height = 22

    header_row(ws4, 3, range(1, 5),
               ["Categorie", "Voorbeeld requirement", "Verborgen architecturele vraag", "Probing Question"])
    asfr_cats = [
        ["Audit Trail",              "Registreer elke wijziging van klantgegevens.",
         "Regulatoir of intern? Hoe lang bewaren?",
         "Is de audit voor interne of externe compliance?"],
        ["Batch Processing",         "Het uitbetalingsproces is een dagelijks batchproces.",
         "Frequentie? Toegangsniveau? Wat bij uitloop?",
         "Waarom batch — technologie of echte functionele behoefte?"],
        ["Localization / Meertalig", "Het systeem ondersteunt meerdere talen.",
         "Welke talen? Rechterschrift? COTS-ondersteuning?",
         "Zijn er regionale wettelijke eisen aan contentlevering?"],
        ["Notificaties / Alerts",    "Stuur een notificatie bij een transactie.",
         "Eenrichtingsverkeer of bevestiging vereist?",
         "Wil de ontvanger een bevestiging terugsturen?"],
        ["Data-invoermechanismen",   "Voer bonuscode in via QR-code of handmatig.",
         "Welke invoerkanalen? Welke externe devices?",
         "Welke standaard-URL-schema's zijn van toepassing?"],
        ["Betaling",                 "Betaling via kaart of internetbankieren.",
         "Één bank of meerdere? Welke betaalprovider?",
         "Hoe wordt succes/falen van de transactie gemeld?"],
        ["Rapportage",               "Genereer maandelijkse klachtenrapporten.",
         "Operationeel of analytisch? Real-time of batch?",
         "Bevat het rapport ongestructureerde documenten?"],
        ["Zoeken",                   "Zoek claimrecords op naam of nummer.",
         "Full-text of gefacetteerd? Verwacht datavolume?",
         "Gedistribueerd of centraal zoekindex?"],
        ["Third Party Interaction",  "Exporteer aanvragen naar backoffice leverancier X.",
         "Binnen of buiten eigen architectuur? Integratiepat roon?",
         "Welke patronen accepteert de derde partij?"],
        ["Workflow",                 "Workflow voor beoordeling en goedkeuring.",
         "Configureerbaar of hard-coded? Escalatiemechanisme?",
         "Zijn workflows lineair of kunnen ze parallel lopen?"],
        ["Online Help",              "Er is een online help-faciliteit beschikbaar.",
         "Consumeren of ook bijdragen? Intern of extern?",
         "Contextgevoelig of generieke zoekfunctie?"],
        ["Licenties",               "Faciliteer bewaking van licentiegebruik.",
         "Gebruikers, CPU-gebaseerd of perpetueel?",
         "Past het licentiemodel bij de COTS-keuze?"],
        ["Gebruikersgedraganalyse",  "Analyseer gebruikersgedrag op het portaal.",
         "Welke metrics? Eerste- of derdepartij tracking?",
         "AVG-impact van de tracking-aanpak?"],
        ["Opslagmechanismen",        "Sla alle documenten op voor x jaar.",
         "Hoe opslaan? Hoe doorzoekbaar? Migratie bij veroudering?",
         "Wat zijn de wettelijke bewaartermijnen?"],
    ]
    for r, row in enumerate(asfr_cats, 4):
        for c, val in enumerate(row, 1):
            cell = ws4.cell(row=r, column=c, value=val)
            cell.font = font(size=9, bold=(c == 1))
            cell.fill = fill(PAPER if r % 2 == 0 else CARD)
            cell.border = thin_border()
            cell.alignment = left()
        ws4.row_dimensions[r].height = 28
    set_col_widths(ws4, [22, 30, 34, 34])
    ws4.freeze_panes = "A4"

    wb.save(OUT / "mva-asr.xlsx")
    print("✓  mva-asr.xlsx")


def build_excel_paved_roads():
    wb = Workbook()
    ws = wb.active
    ws.title = "Paved Road"
    ws.sheet_view.showGridLines = False

    # ── Classificatie kleur-mapping ────────────────────────────────────────────
    classif_fg = {
        "🟢 On-road":     "166534",
        "🟡 Toegestaan":  "854d0e",
        "🔴 Verboden":    "991b1b",
        "📐 Principe":    "1e3a8a",
        "🔍 Kwalificatie":"524d42",
        "📋 Richtlijn":   "524d42",
    }
    classif_bg = {
        "🟢 On-road":     "D0EEDC",
        "🟡 Toegestaan":  "FDF5E0",
        "🔴 Verboden":    "FECACA",
        "📐 Principe":    "DCEEFF",
        "🔍 Kwalificatie":"ECE6D8",
        "📋 Richtlijn":   "F3EFE6",
    }
    layer_bg = {
        1: "EAF5EB", 2: "FFFDE7", 3: "E3F2FD",
        4: "F3E5F5", 5: "FCE4EC", 6: "E8EAF6",
        7: "FBE9E7", 8: "F5F5F5",
    }

    # ── Titel ─────────────────────────────────────────────────────────────────
    ws.merge_cells("A1:J1")
    t = ws["A1"]
    t.value = "Paved Road catalogus — [Organisatie / Programma]"
    t.font = font(bold=True, size=14, color=WHITE)
    t.fill = fill(PRAC)
    t.alignment = center()
    ws.row_dimensions[1].height = 32

    ws.merge_cells("A2:J2")
    s = ws["A2"]
    s.value = ("Gebruik dit template als levend overzichtsdocument van de volledige Paved Road — "
               "alle lagen, kaders en standaarden in één overzicht. "
               "Filter op Laag of Classificatie om te focussen. "
               "MVG-eigenaarschap: principes (Laag 1) worden vastgesteld door MVG en vloeien in naar MVA.")
    s.font = font(size=9, italic=True, color=SOFT)
    s.fill = fill(PAPER2)
    s.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    ws.row_dimensions[2].height = 28

    # ── Metagegevens (links, rijen 3-5) + Status legenda (rechts) ────────────
    for i, (label, val) in enumerate([("Versie", "0.1"), ("Eigenaar", "[rol]"), ("Vastgesteld op", "JJJJ-MM-DD")], 3):
        lc = ws.cell(row=i, column=1, value=label)
        lc.font = font(bold=True, size=9, color=PRAC)
        lc.fill = fill(PAPER2)
        vc = ws.cell(row=i, column=2, value=val)
        vc.font = font(size=9, color=SOFT, italic=True)
        vc.fill = fill(CARD)
        ws.merge_cells(f"B{i}:E{i}")
        ws.row_dimensions[i].height = 16

    # Status legenda (kolommen G-J, rijen 3-8)
    sh = ws.cell(row=3, column=7, value="Status legenda")
    sh.font = font(bold=True, size=9, color=WHITE)
    sh.fill = fill(SOFT)
    sh.alignment = center()
    ws.merge_cells("G3:J3")
    ws.row_dimensions[3].height = 16

    legend = [
        ("🌱 Seeded",    "Eerste opzet, niet gepubliceerd",       "FDF5E0", INK),
        ("📢 Published", "Gepubliceerd, beschikbaar voor teams",  "DCEEFF", INK),
        ("✅ Curated",   "Gevalideerd, aanbevolen",               "D0EEDC", INK),
        ("🏗️ Adopted",  "Actief in gebruik bij teams",           "B8E4C9", INK),
        ("🌐 Pervasive", "De-facto standaard over het landgoed", "0F6E63", WHITE),
        ("🔄 Evolving",  "In herziening",                         "EDE8F8", INK),
    ]
    for i, (status, meaning, bg_color, txt_color) in enumerate(legend, 4):
        c1 = ws.cell(row=i, column=7, value=status)
        c1.font = font(bold=True, size=8, color=txt_color)
        c1.fill = fill(bg_color)
        c1.alignment = mid()
        c1.border = thin_border()
        ws.merge_cells(f"H{i}:J{i}")
        c2 = ws.cell(row=i, column=8, value=meaning)
        c2.font = font(size=8, color=SOFT)
        c2.fill = fill(PAPER if i % 2 == 0 else CARD)
        c2.alignment = mid()
        c2.border = thin_border()
        ws.row_dimensions[i].height = 15

    ws.row_dimensions[10].height = 8   # spacer

    # ── Kolomkoppen ──────────────────────────────────────────────────────────
    HDR = 11
    header_row(ws, HDR, range(1, 11),
               ["Laag", "Domein", "ID", "Item / Naam", "Omschrijving / Consequentie",
                "Classificatie", "Status", "Eigenaar", "SR / ADR", "Geldig t/m"],
               bg=PRAC, row_h=18)

    # ── Data-validatie ────────────────────────────────────────────────────────
    classif_dv = DataValidation(
        type="list",
        formula1='"🟢 On-road,🟡 Toegestaan,🔴 Verboden,📐 Principe,🔍 Kwalificatie,📋 Richtlijn"',
        allow_blank=True
    )
    ws.add_data_validation(classif_dv)

    status_dv = DataValidation(
        type="list",
        formula1='"🌱 Seeded,📢 Published,✅ Curated,🏗️ Adopted,🌐 Pervasive,🔄 Evolving"',
        allow_blank=True
    )
    ws.add_data_validation(status_dv)

    # ── Alle paved road items (laag, domein, id, naam, omschrijving, classificatie, status, eigenaar, sr/adr, geldig) ──
    items = [
        # Laag 1 — Principes (vastgesteld door MVG)
        (1,"Principes","P1","Configureer, customiseer niet",
         "COTS-aanpassingen buiten vendor-API vereisen een ADR met eigenaarschapsverklaring.",
         "📐 Principe","✅ Curated","[Arch. Board]","SR-001",""),
        (1,"Principes","P2","Standaarden boven maatwerk",
         "Waar een open standaard bestaat (OAuth 2.0, OIDC, OpenAPI), is die leidend.",
         "📐 Principe","✅ Curated","[Arch. Board]","SR-002",""),
        (1,"Principes","P3","Data bij de bron",
         "Geen kopiëren van data tenzij expliciet verantwoord; bronsysteem-eigenaar is autoriteit.",
         "📐 Principe","✅ Curated","[Arch. Board]","SR-003",""),
        (1,"Principes","P4","Interoperabiliteit by design",
         "Elk systeem dat data uitwisselt exposeert een gedocumenteerde API. Silo's niet geaccepteerd.",
         "📐 Principe","✅ Curated","[Arch. Board]","SR-004",""),
        (1,"Principes","P5","Exitstrategie vóór adoptie",
         "Elk COTS-product heeft een gedocumenteerde migratieroute voordat het in productie gaat.",
         "📐 Principe","✅ Curated","[Arch. Board]","SR-005",""),
        (1,"Principes","P6","Compliance is architectuureis",
         "Wettelijke verplichtingen als falsifieerbaar criterium gedefinieerd vóór vendorselectie.",
         "📐 Principe","✅ Curated","[Arch. Board]","SR-006",""),
        # Laag 2 — Leverancierskader
        (2,"Leverancierskader","V1","Standaard-authenticatie",
         "Leverancier ondersteunt SAML 2.0 of OIDC; geen proprietary login-flow.",
         "🔍 Kwalificatie","📢 Published","[naam]","",""),
        (2,"Leverancierskader","V2","Gedocumenteerde API",
         "API gedocumenteerd via OpenAPI 3.x of gelijkwaardig; geen vendor-only tooling vereist.",
         "🔍 Kwalificatie","📢 Published","[naam]","",""),
        (2,"Leverancierskader","V3","Dataportabiliteitsgarantie",
         "Export in open formaat gegarandeerd; geen lock-in op data-niveau.",
         "🔍 Kwalificatie","📢 Published","[naam]","",""),
        (2,"Leverancierskader","V4","Supportmodel bij platform-updates",
         "Duidelijk supportmodel voor updates die configuratie raken; geen unilaterale breaking changes.",
         "🔍 Kwalificatie","📢 Published","[naam]","",""),
        (2,"Leverancierskader","V5","Gecertificeerd beveiligingsbeleid",
         "Leverancier heeft gecertificeerd ISMS (bijv. ISO 27001 of gelijkwaardig).",
         "🔍 Kwalificatie","📢 Published","[naam]","",""),
        (2,"Leverancierskader","V6","Breaking change communicatie",
         "Breaking changes minimaal 6 maanden van tevoren gecommuniceerd.",
         "🔍 Kwalificatie","📢 Published","[naam]","",""),
        # Laag 3 — Technologiestack
        (3,"Technologiestack","T1","REST/JSON · OpenAPI 3.x",
         "Standaard voor alle nieuwe koppelingen. Contract-first: spec vóór implementatie.",
         "🟢 On-road","🏗️ Adopted","[naam]","SR-003","[datum]"),
        (3,"Technologiestack","T2","OAuth 2.0 + OpenID Connect",
         "Authenticatie- en autorisatiestandaard voor alle diensten.",
         "🟢 On-road","🏗️ Adopted","[naam]","SR-005","[datum]"),
        (3,"Technologiestack","T3","OCI-compatibele containers",
         "Containerisatie-standaard; geen vendor-lock op container-runtime.",
         "🟢 On-road","🏗️ Adopted","[naam]","SR-003","[datum]"),
        (3,"Technologiestack","T4","Structured logging · centrale aggregatie",
         "Gestandaardiseerd logformaat verplicht; gecentraliseerde observability-stack.",
         "🟢 On-road","🏗️ Adopted","[naam]","SR-003","[datum]"),
        (3,"Technologiestack","T5","SOAP/XML voor legacy-koppelingen",
         "Alleen waar migratie niet haalbaar is; tijdelijk, gedocumenteerd in ADR.",
         "🟡 Toegestaan","📢 Published","[naam]","","[datum]"),
        (3,"Technologiestack","T6","GraphQL",
         "Voor specifieke query-intensieve use cases; motivatie en ADR vereist.",
         "🟡 Toegestaan","🌱 Seeded","[naam]","",""),
        (3,"Technologiestack","T7","Vendor-specifieke databases (met exportgarantie)",
         "Mits dataportabiliteit aantoonbaar is; geen lock-in op data-niveau.",
         "🟡 Toegestaan","📢 Published","[naam]","","[datum]"),
        (3,"Technologiestack","T8","Directe databasekoppelingen tussen applicaties",
         "Creëert onzichtbare afhankelijkheden; altijd via API.",
         "🔴 Verboden","✅ Curated","[naam]","SR-003",""),
        (3,"Technologiestack","T9","Proprietary binaire protocollen (geen open spec)",
         "Maakt onafhankelijke implementatie en migratie onmogelijk.",
         "🔴 Verboden","✅ Curated","[naam]","SR-002",""),
        (3,"Technologiestack","T10","Customisatie van COTS buiten de vendor-API",
         "Blokkeert vendor-updates; vereist eigenaarschapsverklaring + ADR.",
         "🔴 Verboden","✅ Curated","[naam]","SR-001",""),
        # Laag 4 — Datadomeinen
        (4,"Datadomeinen","D1","[Datadomein A]",
         "System of record: [systeem A]. Consumenten raadplegen via API — nooit via database-dump.",
         "📋 Richtlijn","🌱 Seeded","[naam]","",""),
        (4,"Datadomeinen","D2","[Datadomein B]",
         "System of record: [systeem B]. Domeinboundary volgt organisatorische verantwoordelijkheid.",
         "📋 Richtlijn","🌱 Seeded","[naam]","",""),
        (4,"Datadomeinen","D3","[Datadomein C]",
         "System of record: [systeem C]. Twee claimen van SoR voor één domein = architectuurincident.",
         "📋 Richtlijn","🌱 Seeded","[naam]","",""),
        # Laag 5 — Identiteit & SSO
        (5,"Identiteit & SSO","I1","OpenID Connect",
         "Primaire identiteitsstandaard voor alle nieuwe integraties.",
         "🟢 On-road","🏗️ Adopted","[naam]","SR-005","[datum]"),
        (5,"Identiteit & SSO","I2","SAML 2.0",
         "Alleen voor legacy-systemen met gedocumenteerde migratietijdlijn naar OIDC.",
         "🟡 Toegestaan","📢 Published","[naam]","","[datum]"),
        (5,"Identiteit & SSO","I3","Federatieve trust (OIDC/SAML)",
         "B2B-koppelingen met externe partijen mits compatibel met centrale IdP.",
         "🟡 Toegestaan","📢 Published","[naam]","",""),
        (5,"Identiteit & SSO","I4","Eigen identity stores per applicatie",
         "Creëert silo's; onmogelijk te besturen vanuit centrale IdP.",
         "🔴 Verboden","✅ Curated","[naam]","SR-005",""),
        (5,"Identiteit & SSO","I5","Gebruikersnaam + wachtwoord zonder MFA",
         "Niet voor eindgebruikers; beheeraccounts: MFA verplicht.",
         "🔴 Verboden","✅ Curated","[naam]","SR-005",""),
        # Laag 6 — Integratiepatronen
        (6,"Integratiepatronen","G1","REST API (pull, synchroon)",
         "Standaard voor bevragingen en enkelvoudige transacties.",
         "🟢 On-road","🏗️ Adopted","[naam]","SR-006","[datum]"),
        (6,"Integratiepatronen","G2","Event/notificatie (push, asynchroon)",
         "Standaard voor statuswijzigingen en notificaties.",
         "🟢 On-road","🏗️ Adopted","[naam]","SR-006","[datum]"),
        (6,"Integratiepatronen","G3","Bestandsuitwisseling (SFTP / AS4)",
         "Alleen voor legacy en bulk-overdrachten; met migratieplan naar API.",
         "🟡 Toegestaan","📢 Published","[naam]","","[datum]"),
        (6,"Integratiepatronen","G4","ESB / message broker",
         "Voor complexe routering of transformatie mits gemotiveerd.",
         "🟡 Toegestaan","🌱 Seeded","[naam]","",""),
        (6,"Integratiepatronen","G5","Point-to-point databasekoppeling",
         "Creëert onzichtbare afhankelijkheden; altijd via API.",
         "🔴 Verboden","✅ Curated","[naam]","SR-003",""),
        # Laag 7 — Databasekeuzes
        (7,"Databasekeuzes","DB1","Relationele database (open standaard)",
         "Voor transactionele data; elke applicatie heeft zijn eigen schema.",
         "🟢 On-road","🏗️ Adopted","[naam]","","[datum]"),
        (7,"Databasekeuzes","DB2","Documentopslag voor semi-gestructureerde data",
         "Alleen voor use cases waarvoor de structuur functioneel passend is.",
         "🟢 On-road","📢 Published","[naam]","",""),
        (7,"Databasekeuzes","DB3","In-memory stores voor sessie- en cachedata",
         "Alleen voor tijdelijke, niet-kritieke data; geen primaire opslag.",
         "🟡 Toegestaan","📢 Published","[naam]","",""),
        (7,"Databasekeuzes","DB4","Vendor-specifieke database (met exportgarantie)",
         "Mits dataportabiliteit aantoonbaar is.",
         "🟡 Toegestaan","📢 Published","[naam]","","[datum]"),
        (7,"Databasekeuzes","DB5","Gedeelde database als integratiemechanisme",
         "Creëert onzichtbare koppeling tussen applicaties; verboden.",
         "🔴 Verboden","✅ Curated","[naam]","SR-003",""),
        # Laag 8 — Aanvullende dimensies
        (8,"Aanvullende dimensies","A1","Logging & auditability",
         "Logformaat, bewaartermijn en inzagerecht als architectuureis — niet als beheerdetail.",
         "📋 Richtlijn","🌱 Seeded","[naam]","",""),
        (8,"Aanvullende dimensies","A2","Toestemmingsbeheer (consent)",
         "Één authoritative consent store; applicaties raadplegen — implementeren niet zelf.",
         "📋 Richtlijn","🌱 Seeded","[naam]","",""),
        (8,"Aanvullende dimensies","A3","API-versioning en lifecycle",
         "Gedocumenteerde deprecatieprocedure; breaking changes minimaal 6 maanden van tevoren.",
         "📋 Richtlijn","🌱 Seeded","[naam]","",""),
        (8,"Aanvullende dimensies","A4","Leveranciersonboarding",
         "Gestandaardiseerd onboardingproces: kwalificatietest, integratieverificatie, acceptatiecriteria.",
         "📋 Richtlijn","🌱 Seeded","[naam]","",""),
        # Lege invulrij
        ("","","","","","","","","",""),
    ]

    DATA_START = HDR + 1
    for r, row in enumerate(items, DATA_START):
        laag = row[0]
        classif = row[5]
        row_bg = layer_bg.get(laag, CARD)
        for c, val in enumerate(row, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.border = thin_border()
            if c == 1:   # Laag — gecentreerd, vetgedrukt
                cell.font = font(size=9, bold=True, color=PRAC if laag else INK)
                cell.fill = fill(row_bg)
                cell.alignment = center()
            elif c == 6 and classif:   # Classificatie — gekleurde badge
                cell.font = font(size=9, bold=True, color=classif_fg.get(classif, INK))
                cell.fill = fill(classif_bg.get(classif, CARD))
                cell.alignment = mid()
            else:
                cell.font = font(size=9, bold=(c in (3, 4)))
                cell.fill = fill(row_bg)
                cell.alignment = left()
        ws.row_dimensions[r].height = 30
        if classif:
            classif_dv.sqref += f"F{r}:F{r}"
            status_dv.sqref += f"G{r}:G{r}"

    set_col_widths(ws, [6, 20, 7, 26, 44, 18, 16, 12, 10, 12])
    ws.freeze_panes = f"A{DATA_START}"
    ws.auto_filter.ref = f"A{HDR}:J{DATA_START - 1 + len(items)}"

    # Conditionele opmaak op Classificatie-kolom
    for classif_val, bg in classif_bg.items():
        ws.conditional_formatting.add(
            f"F{DATA_START}:F{DATA_START + len(items)}",
            FormulaRule(formula=[f'$F{DATA_START}="{classif_val}"'],
                        fill=PatternFill("solid", fgColor=bg))
        )

    wb.save(OUT / "mva-paved-roads.xlsx")
    print("✓  mva-paved-roads.xlsx")



def build_excel_roadmap():
    wb = Workbook()

    # ── Sheet 1: Roadmap-overzicht ───────────────────────────────────────────
    ws = wb.active
    ws.title = "Roadmap"
    ws.sheet_view.showGridLines = False

    ws.merge_cells("A1:L1")
    t = ws["A1"]
    t.value = "Architecture Roadmap — twee sporen: Configuratie (vendor) & Integratie (architect)"
    t.font = font(bold=True, size=13, color=WHITE)
    t.fill = fill(PRAC)
    t.alignment = center()
    ws.row_dimensions[1].height = 26

    ws.merge_cells("A2:L2")
    s = ws["A2"]
    s.value = ("Alleen Horizon 1 gedetailleerd. Beide sporen starten vanuit de Fit/Gap-analyse "
               "en worden gesynchroniseerd op checkpoints. Governance en change control zijn "
               "randvoorwaarden — zet deze op vóórdat de Fit/Gap-analyse begint.")
    s.font = font(size=8, italic=True, color=SOFT)
    s.fill = fill(PAPER2)
    s.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    ws.row_dimensions[2].height = 28

    # Column headers
    header_row(ws, 3, range(1, 13),
               ["WP #", "Werkpakket", "Spoor", "Horizon", "Team", "Start", "Einde",
                "Alpha-toestand (doel)", "Onomkeerbare grens", "Afhankelijkheden", "ADR's", "Status"],
               bg=PRAC)

    horizon_colors = {
        "Baseline":  PAPER2,
        "Horizon 1": TECH_BG,
        "Horizon 2": APP_BG,
        "Toekomst":  PAPER,
    }
    status_colors = {
        "Gepland":     PAPER,
        "Lopend":      MOT_BG,
        "Afgerond":    TECH_BG,
        "Uitgesteld":  "EEEEEE",
    }

    sample = [
        ["—",    "Baseline (huidige toestand)",          "—",             "Baseline",  "—",            "",      "",      "Architecture → Outlined",                    "—",                           "—",      "—",           "Afgerond"],
        ["WP-01","Platform-fundament",                    "Configuratie",  "Horizon 1", "Platform-team","[Q1]",  "[Q1]",  "Paved Road → Established",                   "—",                           "Budget-goedkeuring", "ADR-01",  "Gepland"],
        ["WP-02","Fit/Gap-analyse & systeemkeuze",        "Configuratie",  "Horizon 1", "Arch-team",   "[Q1]",  "[Q1]",  "Architecture → Outlined\nDecisions → Framed",  "Vendor lock-in · Datamodel",  "WP-01",  "ADR-02",      "Gepland"],
        ["WP-03","Integratie-inventarisatie (vanuit F/G)","Integratie",    "Horizon 1", "Arch-team",   "[Q1]",  "[Q2]",  "Decisions → Framed",                          "—",                           "WP-02",  "ADR-03",      "Gepland"],
        ["WP-04","COTS selectie & PoC",                   "Configuratie",  "Horizon 1", "Arch-team",   "[Q1]",  "[Q2]",  "Architecture → Demonstrated\nDecisions → Decided", "Vendor lock-in definitief", "WP-02", "ADR-02,03,04","Gepland"],
        ["WP-05","Integratieschil (adapters/ACL) fase 1", "Integratie",    "Horizon 1", "Delivery",    "[Q2]",  "[Q2]",  "Paved Road → Adopted (deels)",                "Integratie-API definitief",   "WP-03,04","ADR-03",     "Gepland"],
        ["WP-06","COTS configuratie fase 1",              "Configuratie",  "Horizon 2", "Delivery",    "[Q3]",  "[Q3]",  "System → Demonstrable",                       "—",                           "WP-04",  "—",           "Gepland"],
        ["WP-07","Canoniek datamodel-mapping",            "Integratie",    "Horizon 2", "Arch-team",   "[Q3]",  "[Q4]",  "Paved Road → Adopted",                        "Datamodel definitief",        "WP-05",  "ADR-0x",      "Gepland"],
        ["WP-08","[Fase 2 — nog niet uitgewerkt]",        "—",             "Toekomst",  "[TBD]",       "[H3]",  "[H3]",  "[TBD]",                                       "—",                           "WP-06,07","[TBD]",      "Uitgesteld"],
        ["",     "",                                      "",              "",          "",            "",      "",      "",                                             "",                            "",       "",            ""],
    ]

    status_dv = DataValidation(
        type="list",
        formula1='"Gepland,Lopend,Afgerond,Uitgesteld,Geannuleerd"',
        allow_blank=True
    )
    ws.add_data_validation(status_dv)

    spoor_dv = DataValidation(
        type="list",
        formula1='"Configuratie,Integratie,Beide,—"',
        allow_blank=True
    )
    ws.add_data_validation(spoor_dv)

    horizon_dv = DataValidation(
        type="list",
        formula1='"Baseline,Horizon 1,Horizon 2,Toekomst"',
        allow_blank=True
    )
    ws.add_data_validation(horizon_dv)

    for r, row in enumerate(sample, 4):
        h_bg = horizon_colors.get(row[3], CARD)
        s_bg = status_colors.get(row[11], CARD)
        for c, val in enumerate(row, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.font = font(size=9, bold=(c <= 2))
            cell.fill = fill(h_bg if c <= 10 else s_bg)
            cell.border = thin_border()
            cell.alignment = left()
        ws.row_dimensions[r].height = 30
        status_dv.sqref += f"L{r}:L{r}"
        spoor_dv.sqref += f"C{r}:C{r}"
        horizon_dv.sqref += f"D{r}:D{r}"

    set_col_widths(ws, [7, 28, 14, 12, 14, 7, 7, 28, 26, 18, 10, 12])
    ws.freeze_panes = "A4"
    ws.auto_filter.ref = f"A3:L{3+len(sample)}"

    ws.conditional_formatting.add(
        f"A4:L{3+len(sample)}",
        FormulaRule(formula=['$D4="Horizon 1"'],
                    fill=PatternFill("solid", fgColor=TECH_BG))
    )
    ws.conditional_formatting.add(
        f"A4:L{3+len(sample)}",
        FormulaRule(formula=['$D4="Horizon 2"'],
                    fill=PatternFill("solid", fgColor=APP_BG))
    )

    # ── Sheet 2: Fit/Gap-analyse ─────────────────────────────────────────────
    ws_fg = wb.create_sheet("Fit-Gap")
    ws_fg.sheet_view.showGridLines = False
    ws_fg.merge_cells("A1:G1")
    tfg = ws_fg["A1"]
    tfg.value = "Fit/Gap-analyse — scharnierpunt van configuratiespoor én integratiespoor"
    tfg.font = font(bold=True, size=12, color=WHITE)
    tfg.fill = fill(PRAC)
    tfg.alignment = center()
    ws_fg.row_dimensions[1].height = 24
    ws_fg.merge_cells("A2:G2")
    sfg = ws_fg["A2"]
    sfg.value = ("De fit/gap-analyse bepaalt tegelijk wat de vendor configureert én wat de architect bouwt. "
                 "Classificeer gaps als: Fit | Configuratie | Procesaanpassing | Integratie | Functionele beperking (ADR vereist).")
    sfg.font = font(size=8, italic=True, color=SOFT)
    sfg.fill = fill(PAPER2)
    sfg.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    ws_fg.row_dimensions[2].height = 28

    header_row(ws_fg, 3, range(1, 8),
               ["Req #", "Requirement", "Sommerville-stap", "Classificatie",
                "Spoor", "Beschrijving aanpak", "WP / ADR"])

    class_dv = DataValidation(
        type="list",
        formula1='"Fit (product dekt het),Configuratie (instelbaar),Procesaanpassing (werkwijze wijzigt),'
                 'Integratie (te bouwen koppeling),Functionele beperking (ADR vereist)"',
        allow_blank=True
    )
    ws_fg.add_data_validation(class_dv)

    spoor_fg_dv = DataValidation(
        type="list",
        formula1='"Configuratie,Integratie,Beide,Governance"',
        allow_blank=True
    )
    ws_fg.add_data_validation(spoor_fg_dv)

    fg_sample = [
        ["REQ-01", "Gebruikersauthenticatie via SSO",              "4 · Configuratie", "Fit (product dekt het)",          "Configuratie", "Vendor ondersteunt OIDC — direct configureerbaar", "WP-01"],
        ["REQ-02", "Koppeling met bronsysteem X (REST)",           "4 · Configuratie", "Integratie (te bouwen koppeling)","Integratie",   "Adapter bouwen — OpenAPI contract opstellen",       "WP-03"],
        ["REQ-03", "Rapportage in legacy-formaat Y",               "3 · Procesaanpassing","Procesaanpassing (werkwijze wijzigt)","Configuratie","Proces wijzigen — gebruik vendor-standaardrapport","WP-06"],
        ["REQ-04", "Exporteer naar extern systeem Z (SOAP)",       "4 · Configuratie", "Integratie (te bouwen koppeling)","Integratie",   "SOAP-adapter — tijdelijk, migratiepad vereist",    "WP-03,ADR-0x"],
        ["REQ-05", "Aanpassing kernconfiguratie buiten vendor-API","3 · Procesaanpassing","Functionele beperking (ADR vereist)","Governance","Customisatie buiten vendor-API — ADR vereist",    "ADR-0x"],
        ["REQ-0x", "", "", "", "", "", ""],
    ]
    class_colors = {
        "Fit (product dekt het)":             TECH_BG,
        "Configuratie (instelbaar)":          APP_BG,
        "Procesaanpassing (werkwijze wijzigt)":MOT_BG,
        "Integratie (te bouwen koppeling)":   IMP_BG,
        "Functionele beperking (ADR vereist)":"EEEEEE",
    }
    for r, row in enumerate(fg_sample, 4):
        bg = class_colors.get(row[3], CARD)
        for c, val in enumerate(row, 1):
            cell = ws_fg.cell(row=r, column=c, value=val)
            cell.font = font(size=9, bold=(c == 1))
            cell.fill = fill(bg)
            cell.border = thin_border()
            cell.alignment = left()
        ws_fg.row_dimensions[r].height = 28
        class_dv.sqref += f"D{r}:D{r}"
        spoor_fg_dv.sqref += f"E{r}:E{r}"

    set_col_widths(ws_fg, [9, 28, 18, 28, 14, 32, 16])
    ws_fg.freeze_panes = "A4"
    ws_fg.auto_filter.ref = f"A3:G{3+len(fg_sample)}"

    # ── Sheet 3: Werkpakketten-detail ────────────────────────────────────────
    ws2 = wb.create_sheet("Werkpakketten-detail")
    ws2.sheet_view.showGridLines = False
    ws2.merge_cells("A1:F1")
    t2 = ws2["A1"]
    t2.value = "Werkpakketten — detail (Horizon 1 gedetailleerd)"
    t2.font = font(bold=True, size=12, color=WHITE)
    t2.fill = fill(PRAC)
    t2.alignment = center()
    ws2.row_dimensions[1].height = 22

    for wp_num, wp_name, spoor, deliverables in [
        ("WP-01", "Platform-fundament", "Configuratie", [
            "Cloud Landing Zone ingericht conform paved road",
            "Identity Provider geconfigureerd (SSO)",
            "CI/CD pipeline operationeel",
            "Monitoring & logging platform actief",
        ]),
        ("WP-02", "Fit/Gap-analyse & systeemkeuze", "Configuratie", [
            "Fit/Gap-tabel ingevuld per requirement",
            "Gap-classificaties overeengekomen met stakeholders",
            "RFP uitgestuurd en beoordeeld",
            "Proof-of-Concept uitgevoerd per finalist",
            "ADR-02 t/m ADR-04 besloten — vendor lock-in definitief",
            "Leverancierscontract getekend",
        ]),
        ("WP-03", "Integratie-inventarisatie (vanuit Fit/Gap)", "Integratie", [
            "Koppelingenlijst opgesteld vanuit Fit/Gap-classificatie 'Integratie'",
            "Prioritering van koppelingen per horizon vastgesteld",
            "OpenAPI-contracten opgesteld vóór implementatie (contract-first)",
            "Integratie-API definitief — onomkeerbare grens gepasseerd",
        ]),
    ]:
        r_start = ws2.max_row + 2
        ws2.merge_cells(f"A{r_start}:G{r_start}")
        hdr = ws2.cell(row=r_start, column=1, value=f"{wp_num} — {wp_name}  ·  Spoor: {spoor}")
        hdr.font = font(bold=True, size=11, color=WHITE)
        hdr.fill = fill(PRAC)
        hdr.alignment = mid()
        ws2.row_dimensions[r_start].height = 20

        header_row(ws2, r_start+1, range(1, 8),
                   ["Deliverable", "Eigenaar", "Streefdatum", "Status",
                    "Kwaliteitscriterium", "Onomkeerbare grens?", "ADR"])
        for i, deliv in enumerate(deliverables):
            r = r_start + 2 + i
            row_data = [deliv, "[naam]", "[datum]", "Open", "[criterium]", "Ja / Nee / [omschrijving]", "[ADR-xx]"]
            for c, val in enumerate(row_data, 1):
                cell = ws2.cell(row=r, column=c, value=val)
                cell.font = font(size=9)
                cell.fill = fill(PAPER if i%2==0 else CARD)
                cell.border = thin_border()
                cell.alignment = left()
            ws2.row_dimensions[r].height = 20

    set_col_widths(ws2, [36, 14, 12, 12, 22, 22, 10])
    ws2.freeze_panes = "A3"

    # ── Sheet 3: Hiaten & Afhankelijkheden ───────────────────────────────────
    ws3 = wb.create_sheet("Hiaten")
    ws3.sheet_view.showGridLines = False
    ws3.merge_cells("A1:F1")
    t3 = ws3["A1"]
    t3.value = "Hiaten & Afhankelijkheden"
    t3.font = font(bold=True, size=12, color=WHITE)
    t3.fill = fill(PRAC)
    t3.alignment = center()
    ws3.row_dimensions[1].height = 22

    header_row(ws3, 2, range(1, 7),
               ["Hiaat / Afhankelijkheid", "Type", "Adresseren in", "WP / ADR", "Prioriteit", "Status"])

    gaps = [
        ["Platform niet beschikbaar",  "Technisch",       "WP-01 · H1",  "ADR-01", "Hoog",   "Open"],
        ["Productkeuze nog open",       "Beslissing",      "WP-02 · H1",  "ADR-02", "Hoog",   "Open"],
        ["Budget-goedkeuring",          "Organisatorisch", "WP-01 · H1",  "—",      "Hoog",   "Lopend"],
        ["Datamigratiestrategie",       "Technisch",       "H2",           "ADR-0x", "Middel", "Uitgesteld"],
        ["Organisatieverandering",      "Organisatorisch", "H2–H3",        "—",      "Laag",   "Uitgesteld"],
        ["",                            "",                "",             "",       "",       ""],
    ]
    p_colors = {"Hoog": "FDE8E0", "Middel": "FDF5E0", "Laag": PAPER}
    for r, row in enumerate(gaps, 3):
        bg = p_colors.get(row[4], CARD)
        for c, val in enumerate(row, 1):
            cell = ws3.cell(row=r, column=c, value=val)
            cell.font = font(size=9, bold=(c==1))
            cell.fill = fill(bg)
            cell.border = thin_border()
            cell.alignment = left()
        ws3.row_dimensions[r].height = 20
    set_col_widths(ws3, [28, 18, 16, 12, 12, 12])
    ws3.freeze_panes = "A3"

    # ── Sheet 4: Alpha-Toestanden ────────────────────────────────────────────
    ws4 = wb.create_sheet("Alpha-Toestanden")
    ws4.sheet_view.showGridLines = False
    ws4.merge_cells("A1:E1")
    t4 = ws4["A1"]
    t4.value = "Alpha-Toestanden per Horizon"
    t4.font = font(bold=True, size=12, color=WHITE)
    t4.fill = fill(PRAC)
    t4.alignment = center()
    ws4.row_dimensions[1].height = 22

    header_row(ws4, 2, range(1, 6),
               ["Alpha", "Baseline (nu)", "Horizon 1 (doel)", "Horizon 2 (doel)", "Einddoel"])

    alphas = [
        ["Architectural Drivers", "Identified",        "Quantified",    "Sustained",        "Sustained"],
        ["Architecture",          "Envisioned",         "Demonstrated",  "Usable",           "Established"],
        ["Architecture Decisions","Needed",             "Decided",       "Validated",        "Validated"],
        ["Paved Road",            "Seeded",             "Established",   "Adopted",          "Evolving"],
        ["System",                "—",                  "Outlined",      "Demonstrable",     "Operational"],
    ]
    a_bg = [MOT_BG, APP_BG, APP_BG, TECH_BG, PAPER2]
    for r, (row, bg) in enumerate(zip(alphas, a_bg), 3):
        for c, val in enumerate(row, 1):
            cell = ws4.cell(row=r, column=c, value=val)
            cell.font = font(size=9, bold=(c==1))
            cell.fill = fill(bg)
            cell.border = thin_border()
            cell.alignment = center() if c > 1 else left()
        ws4.row_dimensions[r].height = 22
    set_col_widths(ws4, [24, 20, 20, 20, 20])
    ws4.freeze_panes = "A3"

    wb.save(OUT / "mva-roadmap.xlsx")
    print("✓  mva-roadmap.xlsx")


# ═══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Building MVA work product templates...")
    build_word_architecture_vision()
    build_word_architecture_definition()
    build_word_adr()
    build_word_review_record()
    build_excel_asr()
    build_excel_paved_roads()
    build_excel_roadmap()
    print(f"\nAll templates written to: {OUT}")
